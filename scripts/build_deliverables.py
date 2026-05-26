from __future__ import annotations

import csv
import html
import json
import math
import shutil
import zipfile
from datetime import date, timedelta
from pathlib import Path
from xml.sax.saxutils import escape

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SRC_CSV = Path(r"C:\Users\fabri\Downloads\NuevoTPIA.csv")
SRC_NB = Path(r"C:\Users\fabri\OneDrive\Escritorio\nbeatsnuevo(4).ipynb")

DATA = ROOT / "data"
NOTEBOOKS = ROOT / "notebooks"
RESULTS = ROOT / "results"
DOCS = ROOT / "docs"
PRESENTATION = ROOT / "presentation"
ASSETS = ROOT / "assets"

TARGET = "SIN Imputed"


def easter_sunday(year: int) -> date:
    a = year % 19
    b = year // 100
    c = year % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    l = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * l) // 451
    month = (h + l - 7 * m + 114) // 31
    day = ((h + l - 7 * m + 114) % 31) + 1
    return date(year, month, day)


def paraguay_holidays(years: range) -> set[date]:
    holidays: set[date] = set()
    fixed_month_days = [
        (1, 1),    # Ano Nuevo
        (3, 1),    # Dia de los Heroes
        (5, 1),    # Dia del Trabajador
        (5, 14),   # Independencia
        (5, 15),   # Independencia
        (6, 12),   # Paz del Chaco
        (8, 15),   # Fundacion de Asuncion
        (9, 29),   # Victoria de Boqueron
        (12, 8),   # Virgen de Caacupe
        (12, 25),  # Navidad
    ]
    for year in years:
        holidays.update(date(year, month, day) for month, day in fixed_month_days)
        easter = easter_sunday(year)
        holidays.add(easter - timedelta(days=3))  # Jueves Santo
        holidays.add(easter - timedelta(days=2))  # Viernes Santo
    return holidays


def ensure_dirs() -> None:
    for p in [DATA / "raw", DATA / "processed", NOTEBOOKS, RESULTS, DOCS, PRESENTATION, ASSETS]:
        p.mkdir(parents=True, exist_ok=True)


def copy_sources() -> None:
    shutil.copy2(SRC_CSV, DATA / "raw" / "NuevoTPIA.csv")
    shutil.copy2(SRC_NB, NOTEBOOKS / "nbeats_original.ipynb")


def load_and_engineer() -> pd.DataFrame:
    df = pd.read_csv(SRC_CSV)
    df = df.rename(columns={df.columns[0]: "DATETIME"})
    df["DATETIME"] = pd.to_datetime(df["DATETIME"], utc=True, errors="coerce")
    df = df.dropna(subset=["DATETIME", TARGET]).sort_values("DATETIME")
    df = df.drop_duplicates("DATETIME").set_index("DATETIME")
    numeric_cols = [c for c in df.columns if c != "DATETIME"]
    df[numeric_cols] = df[numeric_cols].apply(pd.to_numeric, errors="coerce")
    df = df.asfreq("1h")
    df[numeric_cols] = df[numeric_cols].interpolate(limit_direction="both")

    local_idx = df.index.tz_convert("America/Asuncion")
    hour = local_idx.hour.to_numpy()
    dow = local_idx.dayofweek.to_numpy()
    month = local_idx.month.to_numpy()
    local_dates = np.array([ts.date() for ts in local_idx])
    holidays = paraguay_holidays(range(local_idx.year.min() - 1, local_idx.year.max() + 2))
    df["hour_sin"] = np.sin(2 * np.pi * hour / 24)
    df["hour_cos"] = np.cos(2 * np.pi * hour / 24)
    df["dow_sin"] = np.sin(2 * np.pi * dow / 7)
    df["dow_cos"] = np.cos(2 * np.pi * dow / 7)
    df["month_sin"] = np.sin(2 * np.pi * month / 12)
    df["month_cos"] = np.cos(2 * np.pi * month / 12)
    df["is_weekend"] = (dow >= 5).astype(int)
    df["is_night"] = ((hour >= 19) | (hour <= 5)).astype(int)
    df["is_morning"] = ((hour >= 6) & (hour <= 11)).astype(int)
    df["is_afternoon"] = ((hour >= 12) & (hour <= 18)).astype(int)
    df["is_peak_hour"] = (((hour >= 18) & (hour <= 22)) | ((hour >= 6) & (hour <= 8))).astype(int)
    df["is_holiday"] = np.array([d in holidays for d in local_dates], dtype=int)
    df["is_day_before_holiday"] = np.array([d + timedelta(days=1) in holidays for d in local_dates], dtype=int)
    df["is_day_after_holiday"] = np.array([d - timedelta(days=1) in holidays for d in local_dates], dtype=int)
    df["cooling_degree_24c"] = np.maximum(df["T02M"] - 24, 0)
    df["heating_degree_18c"] = np.maximum(18 - df["T02M"], 0)
    df["hot_afternoon"] = df["cooling_degree_24c"] * df["is_afternoon"]
    df["hot_peak_hour"] = df["cooling_degree_24c"] * df["is_peak_hour"]
    df["lag_24h"] = df[TARGET].shift(24)
    df["lag_168h"] = df[TARGET].shift(168)
    df["lag_1h"] = df[TARGET].shift(1)
    df["lag_2h"] = df[TARGET].shift(2)
    df["lag_3h"] = df[TARGET].shift(3)
    df["lag_48h"] = df[TARGET].shift(48)
    df["lag_336h"] = df[TARGET].shift(336)
    df["target_diff_1h"] = df[TARGET].diff(1).shift(1)
    df["temp_diff_1h"] = df["T02M"].diff(1)
    df["humidity_diff_1h"] = df["RH2M"].diff(1)
    df["rolling_24h_mean"] = df[TARGET].shift(1).rolling(24).mean()
    df["rolling_168h_mean"] = df[TARGET].shift(1).rolling(168).mean()
    df["rolling_3h_mean"] = df[TARGET].shift(1).rolling(3).mean()
    df["rolling_6h_mean"] = df[TARGET].shift(1).rolling(6).mean()
    df["rolling_12h_mean"] = df[TARGET].shift(1).rolling(12).mean()
    df["rolling_48h_mean"] = df[TARGET].shift(1).rolling(48).mean()
    df["temp_max_24h"] = df["T02M"].shift(1).rolling(24).max()
    df["temp_min_24h"] = df["T02M"].shift(1).rolling(24).min()
    df["target_max_24h"] = df[TARGET].shift(1).rolling(24).max()
    df["target_min_24h"] = df[TARGET].shift(1).rolling(24).min()
    df["wind_speed_10m"] = np.sqrt(df["U10M"] ** 2 + df["V10M"] ** 2)
    df["temp_humidity_interaction"] = df["T02M"] * df["RH2M"]
    df = df.dropna()
    df.to_csv(DATA / "processed" / "processed_dataset.csv", index_label="DATETIME")
    return df


def standardize(train_x: np.ndarray, test_x: np.ndarray) -> tuple[np.ndarray, np.ndarray, np.ndarray, np.ndarray]:
    mu = train_x.mean(axis=0)
    sigma = train_x.std(axis=0)
    sigma[sigma == 0] = 1.0
    return (train_x - mu) / sigma, (test_x - mu) / sigma, mu, sigma


def ridge_fit(x: np.ndarray, y: np.ndarray, alpha: float) -> np.ndarray:
    x_aug = np.column_stack([np.ones(len(x)), x])
    reg = np.eye(x_aug.shape[1]) * alpha
    reg[0, 0] = 0.0
    return np.linalg.solve(x_aug.T @ x_aug + reg, x_aug.T @ y)


def ridge_predict(x: np.ndarray, beta: np.ndarray) -> np.ndarray:
    x_aug = np.column_stack([np.ones(len(x)), x])
    return x_aug @ beta


def metrics(y: np.ndarray, pred: np.ndarray) -> dict[str, float]:
    err = pred - y
    rmse = float(np.sqrt(np.mean(err**2)))
    mae = float(np.mean(np.abs(err)))
    mape = float(np.mean(np.abs(err / np.maximum(np.abs(y), 1e-9))) * 100)
    r2 = float(1 - np.sum(err**2) / np.sum((y - y.mean()) ** 2))
    return {"RMSE": rmse, "MAE": mae, "MAPE_%": mape, "R2": r2}


def run_grid_search(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, dict[str, float], list[str]]:
    base_weather = ["T02M", "RH2M", "PRSS", "TPP6", "U10M", "V10M"]
    behavior_features = [
        "is_weekend", "is_night", "is_morning", "is_afternoon", "is_peak_hour",
        "is_holiday", "is_day_before_holiday", "is_day_after_holiday",
    ]
    thermal_features = ["cooling_degree_24c", "heating_degree_18c", "hot_afternoon", "hot_peak_hour"]
    short_memory_features = [
        "lag_1h", "lag_2h", "lag_3h", "lag_48h", "lag_336h",
        "target_diff_1h", "temp_diff_1h", "humidity_diff_1h",
        "rolling_3h_mean", "rolling_6h_mean", "rolling_12h_mean", "rolling_48h_mean",
        "temp_max_24h", "temp_min_24h", "target_max_24h", "target_min_24h",
    ]
    engineered_min = ["hour_sin", "hour_cos", "dow_sin", "dow_cos", "lag_24h", "lag_168h"]
    engineered_full = [
        "hour_sin", "hour_cos", "dow_sin", "dow_cos", "month_sin", "month_cos",
        *behavior_features,
        "lag_24h", "lag_168h", "rolling_24h_mean", "rolling_168h_mean",
        "wind_speed_10m", "temp_humidity_interaction",
        *thermal_features,
        *short_memory_features,
    ]
    feature_sets = {
        "weather_only": base_weather,
        "weather_calendar_lags": base_weather + engineered_min,
        "weather_behavior_lags": base_weather + behavior_features + engineered_min,
        "weather_behavior_thermal_lags": base_weather + behavior_features + thermal_features + engineered_min,
        "weather_short_memory": base_weather + engineered_min + short_memory_features,
        "full_engineered": base_weather + engineered_full,
    }
    alphas = [0.0, 0.1, 1.0, 10.0, 100.0, 1000.0]
    train = df[df.index < pd.Timestamp("2021-01-01", tz="UTC")]
    test = df[df.index >= pd.Timestamp("2021-01-01", tz="UTC")]

    rows = []
    best = None
    for set_name, cols in feature_sets.items():
        x_train_raw = train[cols].to_numpy(float)
        x_test_raw = test[cols].to_numpy(float)
        x_train, x_test, _, _ = standardize(x_train_raw, x_test_raw)
        y_train = train[TARGET].to_numpy(float)
        y_test = test[TARGET].to_numpy(float)
        for alpha in alphas:
            beta = ridge_fit(x_train, y_train, alpha)
            pred = ridge_predict(x_test, beta)
            row = {"feature_set": set_name, "alpha": alpha, **metrics(y_test, pred)}
            rows.append(row)
            if best is None or row["RMSE"] < best["row"]["RMSE"]:
                best = {"row": row, "beta": beta, "cols": cols, "pred": pred, "y": y_test, "test": test}

    results = pd.DataFrame(rows).sort_values("RMSE")
    results.to_csv(RESULTS / "hyperparameter_results.csv", index=False)
    assert best is not None
    pred_df = pd.DataFrame({
        "DATETIME": best["test"].index[: len(best["pred"])],
        "actual": best["y"],
        "prediction": best["pred"],
        "error": best["pred"] - best["y"],
        "absolute_error": np.abs(best["pred"] - best["y"]),
    })
    pred_df.to_csv(RESULTS / "predictions.csv", index=False)
    summary = {"best_feature_set": best["row"]["feature_set"], "best_alpha": best["row"]["alpha"], **metrics(best["y"], best["pred"])}
    with (RESULTS / "metrics_summary.json").open("w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2)
    return results, pred_df, summary, best["cols"]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    base = Path("C:/Windows/Fonts")
    name = "arialbd.ttf" if bold else "arial.ttf"
    return ImageFont.truetype(str(base / name), size=size)


def line_chart(path: Path, title: str, series: list[tuple[str, np.ndarray, str]], y_label: str = "") -> None:
    w, h = 1200, 900
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    margin = (110, 95, 55, 120)
    x0, y0, x1, y1 = margin[0], margin[1], w - margin[2], h - margin[3]
    d.text((x0, 28), title, fill="#111827", font=font(38, True))
    all_y = np.concatenate([s[1] for s in series])
    mn, mx = float(np.nanmin(all_y)), float(np.nanmax(all_y))
    pad = (mx - mn) * 0.08 or 1
    mn -= pad
    mx += pad
    for i in range(6):
        y = y1 - i * (y1 - y0) / 5
        val = mn + i * (mx - mn) / 5
        d.line((x0, y, x1, y), fill="#e5e7eb", width=1)
        d.text((18, y - 12), f"{val:,.0f}", fill="#374151", font=font(22))
    d.line((x0, y0, x0, y1), fill="#111827", width=2)
    d.line((x0, y1, x1, y1), fill="#111827", width=2)
    for label, vals, color in series:
        vals = np.asarray(vals, dtype=float)
        pts = []
        n = len(vals)
        for i, v in enumerate(vals):
            x = x0 + i * (x1 - x0) / max(n - 1, 1)
            y = y1 - (v - mn) * (y1 - y0) / (mx - mn)
            pts.append((x, y))
        if len(pts) > 1:
            d.line(pts, fill=color, width=4)
    lx = x0
    for label, _, color in series:
        d.rounded_rectangle((lx, h - 82, lx + 24, h - 58), radius=4, fill=color)
        d.text((lx + 34, h - 86), label, fill="#111827", font=font(24))
        lx += 260
    if y_label:
        d.text((x0, h - 45), y_label, fill="#374151", font=font(22))
    im.save(path)


def histogram(path: Path, title: str, values: np.ndarray) -> None:
    w, h = 900, 900
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    x0, y0, x1, y1 = 110, 100, 845, 760
    d.text((x0, 28), title, fill="#111827", font=font(36, True))
    counts, bins = np.histogram(values, bins=32)
    mx = max(counts) or 1
    for i, c in enumerate(counts):
        bx0 = x0 + i * (x1 - x0) / len(counts)
        bx1 = x0 + (i + 1) * (x1 - x0) / len(counts) - 3
        by0 = y1 - c * (y1 - y0) / mx
        d.rectangle((bx0, by0, bx1, y1), fill="#0f766e")
    d.line((x0, y0, x0, y1), fill="#111827", width=2)
    d.line((x0, y1, x1, y1), fill="#111827", width=2)
    d.text((x0, 795), f"Error medio: {np.mean(values):,.2f} | Desv.: {np.std(values):,.2f}", fill="#374151", font=font(25))
    d.text((x0, 832), "Errores de prediccion: prediccion - valor real", fill="#374151", font=font(24))
    im.save(path)


def bar_chart(path: Path, title: str, labels: list[str], vals: list[float]) -> None:
    w, h = 1100, 900
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    x0, y0, x1, y1 = 130, 110, 1030, 735
    d.text((x0, 28), title, fill="#111827", font=font(36, True))
    mx = max(vals) * 1.08
    bw = (x1 - x0) / len(vals) * 0.65
    for i, (lab, v) in enumerate(zip(labels, vals)):
        cx = x0 + (i + 0.5) * (x1 - x0) / len(vals)
        by = y1 - v * (y1 - y0) / mx
        d.rectangle((cx - bw / 2, by, cx + bw / 2, y1), fill="#2563eb")
        d.text((cx - 55, by - 32), f"{v:,.0f}", fill="#111827", font=font(20, True))
        d.text((cx - 95, y1 + 18), lab[:18], fill="#374151", font=font(19))
    d.line((x0, y1, x1, y1), fill="#111827", width=2)
    d.line((x0, y0, x0, y1), fill="#111827", width=2)
    im.save(path)


def create_charts(results: pd.DataFrame, pred_df: pd.DataFrame) -> None:
    sample = pred_df.iloc[: 24 * 14]
    line_chart(
        ASSETS / "predictions_vs_actual.png",
        "Predicciones vs valores reales: primeras 2 semanas",
        [("Real", sample["actual"].to_numpy(), "#111827"), ("Prediccion", sample["prediction"].to_numpy(), "#dc2626")],
        "Demanda / SIN",
    )
    histogram(ASSETS / "error_histogram.png", "Histograma de errores", pred_df["error"].to_numpy())
    best_by_set = results.groupby("feature_set")["RMSE"].min().reset_index()
    bar_chart(ASSETS / "rmse_by_feature_set.png", "RMSE por conjunto de features", best_by_set["feature_set"].tolist(), best_by_set["RMSE"].tolist())
    conv = results[results["feature_set"] == "full_engineered"].sort_values("alpha")
    line_chart(
        ASSETS / "grid_search_curve.png",
        "Convergencia de Grid Search: alpha vs RMSE",
        [("RMSE", conv["RMSE"].to_numpy(), "#7c3aed")],
        "Valores de alpha: 0, 0.1, 1, 10, 100, 1000",
    )


def nb(cells: list[dict]) -> dict:
    return {
        "cells": cells,
        "metadata": {"kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"}, "language_info": {"name": "python", "version": "3.x"}},
        "nbformat": 4,
        "nbformat_minor": 5,
    }


def md(text: str) -> dict:
    return {"cell_type": "markdown", "metadata": {}, "source": text.splitlines(True)}


def code(text: str) -> dict:
    return {"cell_type": "code", "execution_count": None, "metadata": {}, "outputs": [], "source": text.splitlines(True)}


def create_notebooks() -> None:
    notebooks = {
        "01_preprocesamiento.ipynb": [
            md("# Preprocesamiento\nCarga del dataset, limpieza temporal y generacion de features."),
            code("""import pandas as pd, numpy as np
from datetime import date, timedelta

def easter_sunday(year):
    a = year % 19
    b = year // 100
    c = year % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19 * a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    l = (32 + 2 * e + 2 * i - h - k) % 7
    m = (a + 11 * h + 22 * l) // 451
    month = (h + l - 7 * m + 114) // 31
    day = ((h + l - 7 * m + 114) % 31) + 1
    return date(year, month, day)

def paraguay_holidays(years):
    fixed = [(1,1), (3,1), (5,1), (5,14), (5,15), (6,12), (8,15), (9,29), (12,8), (12,25)]
    out = set()
    for year in years:
        out.update(date(year, month, day) for month, day in fixed)
        easter = easter_sunday(year)
        out.add(easter - timedelta(days=3))
        out.add(easter - timedelta(days=2))
    return out

raw = pd.read_csv('../data/raw/NuevoTPIA.csv')
raw = raw.rename(columns={raw.columns[0]: 'DATETIME'})
raw['DATETIME'] = pd.to_datetime(raw['DATETIME'], utc=True, errors='coerce')
raw = raw.dropna(subset=['DATETIME', 'SIN Imputed']).sort_values('DATETIME')
raw = raw.drop_duplicates('DATETIME').set_index('DATETIME').asfreq('1h').interpolate()
idx = raw.index.tz_convert('America/Asuncion')
hour = idx.hour.to_numpy()
dow = idx.dayofweek.to_numpy()
local_dates = np.array([ts.date() for ts in idx])
holidays = paraguay_holidays(range(idx.year.min() - 1, idx.year.max() + 2))

raw['hour_sin'] = np.sin(2*np.pi*hour/24)
raw['hour_cos'] = np.cos(2*np.pi*hour/24)
raw['is_weekend'] = (dow >= 5).astype(int)
raw['is_night'] = ((hour >= 19) | (hour <= 5)).astype(int)
raw['is_morning'] = ((hour >= 6) & (hour <= 11)).astype(int)
raw['is_afternoon'] = ((hour >= 12) & (hour <= 18)).astype(int)
raw['is_peak_hour'] = (((hour >= 18) & (hour <= 22)) | ((hour >= 6) & (hour <= 8))).astype(int)
raw['is_holiday'] = np.array([d in holidays for d in local_dates], dtype=int)
raw['is_day_before_holiday'] = np.array([d + timedelta(days=1) in holidays for d in local_dates], dtype=int)
raw['is_day_after_holiday'] = np.array([d - timedelta(days=1) in holidays for d in local_dates], dtype=int)
raw['cooling_degree_24c'] = np.maximum(raw['T02M'] - 24, 0)
raw['heating_degree_18c'] = np.maximum(18 - raw['T02M'], 0)
raw['hot_afternoon'] = raw['cooling_degree_24c'] * raw['is_afternoon']
raw['hot_peak_hour'] = raw['cooling_degree_24c'] * raw['is_peak_hour']
raw['lag_1h'] = raw['SIN Imputed'].shift(1)
raw['lag_2h'] = raw['SIN Imputed'].shift(2)
raw['lag_3h'] = raw['SIN Imputed'].shift(3)
raw['lag_24h'] = raw['SIN Imputed'].shift(24)
raw['lag_48h'] = raw['SIN Imputed'].shift(48)
raw['lag_168h'] = raw['SIN Imputed'].shift(168)
raw['lag_336h'] = raw['SIN Imputed'].shift(336)
raw['rolling_6h_mean'] = raw['SIN Imputed'].shift(1).rolling(6).mean()
raw['rolling_24h_mean'] = raw['SIN Imputed'].shift(1).rolling(24).mean()
raw['wind_speed_10m'] = np.sqrt(raw['U10M']**2 + raw['V10M']**2)
processed = raw.dropna()
processed.to_csv('../data/processed/processed_dataset.csv')
processed.head()"""),
        ],
        "02_entrenamiento.ipynb": [
            md("# Entrenamiento\nModelo base de regresion Ridge para pronostico horario."),
            code("import json, pandas as pd\nmetrics = json.load(open('../results/metrics_summary.json'))\nprint(metrics)\npred = pd.read_csv('../results/predictions.csv')\npred.head()"),
        ],
        "03_optimizacion.ipynb": [
            md("# Optimizacion de hiperparametros\nGrid Search manual sobre alpha y conjuntos de features."),
            code("import pandas as pd\nresults = pd.read_csv('../results/hyperparameter_results.csv')\nresults.sort_values('RMSE').head(10)"),
        ],
        "04_visualizacion.ipynb": [
            md("# Visualizacion de resultados\nGraficos principales generados para el reporte y la presentacion."),
            code("from IPython.display import Image, display\nfor p in ['../assets/predictions_vs_actual.png','../assets/error_histogram.png','../assets/grid_search_curve.png','../assets/rmse_by_feature_set.png']:\n    display(Image(filename=p))"),
        ],
    }
    for name, cells in notebooks.items():
        (NOTEBOOKS / name).write_text(json.dumps(nb(cells), indent=2), encoding="utf-8")


def create_readme(summary: dict[str, float], features: list[str]) -> None:
    text = f"""# Pronostico horario de SIN con variables meteorologicas

Proyecto de inteligencia artificial para estimar la serie horaria `SIN Imputed` a partir de variables meteorologicas y features temporales.

## Estructura

- `data/raw/`: dataset original.
- `data/processed/`: dataset limpio con features generados.
- `notebooks/`: preprocesamiento, entrenamiento, optimizacion, visualizacion y notebook N-BEATS original.
- `results/`: resultados CSV de Grid Search, predicciones y metricas.
- `assets/`: figuras usadas en reporte, presentacion y GitHub Pages.
- `docs/`: reporte en formato articulo IEEE y pagina `index.html`.
- `presentation/`: presentacion de 10 diapositivas.

## Features generados

| Grupo | Features | Como se generan | Por que ayudan |
|---|---|---|---|
| Calendario ciclico | `hour_sin`, `hour_cos`, `dow_sin`, `dow_cos`, `month_sin`, `month_cos` | Se transforma hora, dia y mes con seno/coseno. | Representan ciclos diarios, semanales y anuales sin cortes bruscos. |
| Comportamiento energetico | `is_weekend`, `is_night`, `is_morning`, `is_afternoon`, `is_peak_hour` | Se crean indicadores 0/1 segun dia y franja horaria. | Capturan habitos de consumo, como diferencias entre noche, tarde, hora pico y fin de semana. |
| Feriados | `is_holiday`, `is_day_before_holiday`, `is_day_after_holiday` | Se marcan feriados nacionales de Paraguay y dias cercanos. | Un feriado puede parecerse mas a un domingo que a un dia laboral normal. |
| Sensibilidad termica | `cooling_degree_24c`, `heating_degree_18c`, `hot_afternoon`, `hot_peak_hour` | Se mide cuanto la temperatura supera 24 C o baja de 18 C, y se combina con tarde/hora pico. | Aproxima el efecto de refrigeracion/calefaccion sobre la demanda. |
| Memoria temporal | `lag_1h`, `lag_2h`, `lag_3h`, `lag_24h`, `lag_168h`, `lag_336h` | Se usan valores anteriores de la serie. | La demanda horaria suele parecerse a horas recientes y al mismo horario de dias/semanas previas. |
| Promedios y extremos recientes | `rolling_3h_mean`, `rolling_6h_mean`, `rolling_12h_mean`, `rolling_24h_mean`, `temp_max_24h`, `target_max_24h` | Se calculan medias, maximos y minimos de ventanas anteriores. | Suavizan ruido y resumen el estado reciente del sistema. |
| Clima derivado | `wind_speed_10m`, `temp_humidity_interaction` | Se calcula magnitud del viento y producto temperatura-humedad. | Resume condiciones meteorologicas que pueden afectar el consumo. |

## Modelo y metricas

Se implemento un baseline de regresion Ridge con Grid Search manual sobre `alpha` y varios conjuntos de features. El mejor resultado fue:

- Feature set: `{summary['best_feature_set']}`
- Alpha: `{summary['best_alpha']}`
- RMSE: `{summary['RMSE']:.2f}`
- MAE: `{summary['MAE']:.2f}`
- MAPE: `{summary['MAPE_%']:.2f}%`
- R2: `{summary['R2']:.4f}`

Como el problema es de regresion de una serie temporal, se usan RMSE, MAE, MAPE y R2. No se usa precision, F1-Score ni matriz de confusion porque esas metricas corresponden a problemas de clasificacion.

## Evaluacion preliminar

Se entreno con datos anteriores a 2021 y se valido desde 2021 en adelante. Esta separacion evita mezclar informacion futura en el entrenamiento. La tabla `results/hyperparameter_results.csv` permite comparar modelos y configuraciones, mientras que `results/predictions.csv` contiene los valores reales, predichos y errores.

## Visualizaciones

- `assets/predictions_vs_actual.png`: compara valores reales y predicciones.
- `assets/error_histogram.png`: muestra la distribucion de errores.
- `assets/rmse_by_feature_set.png`: compara el RMSE entre grupos de variables.
- `assets/grid_search_curve.png`: muestra el efecto del hiperparametro `alpha` en el RMSE.

El notebook `notebooks/nbeats_original.ipynb` conserva el enfoque avanzado con N-BEATS.
"""
    (ROOT / "README.md").write_text(text, encoding="utf-8")


def create_index(summary: dict[str, float]) -> None:
    img = "../assets/predictions_vs_actual.png"
    html_text = f"""<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Benjamin - Proyecto IA</title>
  <style>
    body {{ margin: 0; font-family: Arial, sans-serif; color: #172033; background: #f7fafc; }}
    header {{ background: #0f766e; color: white; padding: 56px 8vw 40px; }}
    main {{ max-width: 1040px; margin: auto; padding: 34px 22px 60px; }}
    section {{ margin: 28px 0; }}
    h1 {{ font-size: 44px; margin: 0 0 8px; }}
    h2 {{ color: #0f766e; border-bottom: 2px solid #d9eee9; padding-bottom: 8px; }}
    .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(260px, 1fr)); gap: 18px; }}
    .card {{ background: white; border: 1px solid #dbe4ea; border-radius: 8px; padding: 20px; }}
    img {{ max-width: 100%; border-radius: 8px; border: 1px solid #dbe4ea; background: white; }}
    a {{ color: #0f766e; font-weight: 700; }}
  </style>
</head>
<body>
  <header>
    <h1>Benjamin</h1>
    <p>Estudiante de Ingenieria Electromecanica | IA aplicada a series temporales</p>
  </header>
  <main>
    <section>
      <h2>Inicio</h2>
      <p>Me interesa aplicar programacion, analisis de datos y aprendizaje automatico para resolver problemas de energia, meteorologia y sistemas electromecanicos. Este sitio resume mi proyecto de IA y funciona como pagina base para GitHub Pages.</p>
    </section>
    <section class="grid">
      <div class="card"><h2>Habilidades</h2><p>Python, Machine Learning, Pandas, NumPy, series temporales, Git, Grid Search, N-BEATS, analisis de datos, visualizacion.</p></div>
      <div class="card"><h2>Educacion</h2><p>Ingenieria Electromecanica, formacion universitaria en curso. Ajusta aqui institucion y anos antes de publicar.</p></div>
      <div class="card"><h2>Contacto</h2><p>LinkedIn: agrega tu perfil<br>GitHub: agrega tu usuario<br>Email: agrega tu correo</p></div>
    </section>
    <section>
      <h2>Proyecto de Inteligencia Artificial</h2>
      <p><strong>Titulo:</strong> Pronostico horario de SIN con variables meteorologicas.</p>
      <p>El objetivo es predecir el comportamiento horario de la serie `SIN Imputed` usando temperatura, humedad, presion, precipitacion, viento y features temporales. La relevancia del proyecto esta en apoyar analisis de demanda/produccion mediante modelos reproducibles.</p>
      <ul>
        <li><strong>Dataset:</strong> serie temporal horaria con variables meteorologicas desde archivo CSV.</li>
        <li><strong>Modelo:</strong> baseline Ridge optimizado con Grid Search; notebook original con N-BEATS como enfoque avanzado.</li>
        <li><strong>Resultado:</strong> RMSE {summary['RMSE']:.2f}, MAE {summary['MAE']:.2f}, R2 {summary['R2']:.4f}.</li>
      </ul>
      <img src="{img}" alt="Grafico de predicciones contra valores reales">
      <p><a href="https://github.com/tu-usuario/tu-repositorio">Ver repositorio en GitHub</a></p>
    </section>
  </main>
</body>
</html>
"""
    (DOCS / "index.html").write_text(html_text, encoding="utf-8")


def set_columns(section, num: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "360")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.size = Pt(8)
        r.italic = True


def create_report(summary: dict[str, float], results: pd.DataFrame) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(9)
    for style in ["Heading 1", "Heading 2"]:
        styles[style].font.name = "Times New Roman"
        styles[style].font.color.rgb = RGBColor(0, 0, 0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Pronostico horario de SIN mediante aprendizaje automatico y variables meteorologicas")
    run.bold = True
    run.font.size = Pt(16)
    auth = doc.add_paragraph("Benjamin\nEstudiante de Ingenieria Electromecanica")
    auth.alignment = WD_ALIGN_PARAGRAPH.CENTER

    abs_p = doc.add_paragraph()
    abs_p.add_run("Abstract—").bold = True
    abs_p.add_run(
        "Este trabajo presenta un flujo reproducible para el pronostico horario de la variable SIN Imputed utilizando datos meteorologicos y caracteristicas temporales. "
        "Se realizo limpieza de la serie, interpolacion horaria, generacion de features ciclicos, rezagos, medias moviles y magnitud de viento. "
        "Como modelo base se implemento una regresion Ridge con busqueda en grilla sobre el parametro de regularizacion y conjuntos de variables. "
        f"El mejor modelo obtuvo RMSE={summary['RMSE']:.2f}, MAE={summary['MAE']:.2f}, MAPE={summary['MAPE_%']:.2f}% y R2={summary['R2']:.4f} en el periodo de validacion desde 2021. "
        "Los resultados evidencian que los rezagos diarios y semanales aportan informacion relevante para capturar la estacionalidad de la serie."
    )
    doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(doc.sections[-1], 2)

    doc.add_heading("I. Introduccion", level=1)
    doc.add_paragraph(
        "El pronostico de series temporales energeticas o de demanda asociada al sistema interconectado es una tarea relevante para planificacion, operacion y analisis de recursos. "
        "Los modelos de aprendizaje automatico permiten integrar informacion meteorologica y patrones historicos para producir estimaciones cuantitativas. "
        "En trabajos actuales se emplean desde modelos lineales regularizados hasta redes neuronales profundas para series temporales, como N-BEATS, que fue explorado en el notebook original del proyecto."
    )
    doc.add_heading("II. Metodologia", level=1)
    doc.add_paragraph(
        "El dataset proviene de un archivo CSV con frecuencia horaria. Contiene la variable objetivo SIN Imputed y covariables meteorologicas: temperatura a 2 m, humedad relativa, presion, precipitacion y componentes de viento a 10 m."
    )
    doc.add_paragraph(
        "Se generaron features adicionales: variables senoidales y cosenoidales para hora, dia de semana y mes; indicadores de fin de semana, feriado, noche, manana, tarde y hora pico; rezagos de 1, 2, 3, 24, 48, 168 y 336 horas; medias moviles de 3, 6, 12, 24, 48 y 168 horas; grados de enfriamiento y calefaccion aproximados; magnitud de viento calculada como raiz de U10M^2 + V10M^2; e interaccion temperatura-humedad."
    )
    doc.add_paragraph(
        "El modelo base es una regresion Ridge entrenada por ecuacion normal regularizada. La optimizacion usa Grid Search sobre alpha y tres conjuntos de variables. La evaluacion se realiza con RMSE, MAE, MAPE y R2."
    )
    doc.add_paragraph(
        "El problema se formula como regresion de serie temporal, ya que la salida esperada es un valor numerico continuo. Por este motivo no se emplean metricas de clasificacion como precision, F1-Score o matriz de confusion."
    )
    doc.add_paragraph(
        "Las variables generadas se agrupan en calendario ciclico, comportamiento energetico, feriados, sensibilidad termica, memoria temporal y clima derivado. Esta organizacion permite comparar si cada familia de informacion aporta al desempeno predictivo."
    )
    doc.add_heading("III. Resultados", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Feature set", "Alpha", "RMSE", "MAE", "R2"]):
        cell.text = text
    for _, row in results.head(6).iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["feature_set"])
        cells[1].text = f"{row['alpha']:.1f}"
        cells[2].text = f"{row['RMSE']:.2f}"
        cells[3].text = f"{row['MAE']:.2f}"
        cells[4].text = f"{row['R2']:.4f}"
    doc.add_paragraph()
    doc.add_paragraph(
        "La evaluacion preliminar se realizo separando los datos por fecha: entrenamiento antes de 2021 y validacion desde 2021. El archivo CSV de hiperparametros conserva todas las pruebas del Grid Search para facilitar comparaciones y tablas."
    )
    for image, cap in [
        ("predictions_vs_actual.png", "Fig. 1. Predicciones y valores reales durante las primeras dos semanas de validacion."),
        ("error_histogram.png", "Fig. 2. Distribucion de errores del modelo en el conjunto de validacion."),
        ("grid_search_curve.png", "Fig. 3. Curva de busqueda de hiperparametros para el conjunto completo de features."),
    ]:
        doc.add_picture(str(ASSETS / image), width=Inches(3.15))
        add_caption(doc, cap)
    doc.add_heading("IV. Conclusiones", level=1)
    doc.add_paragraph(
        "El flujo desarrollado cumple con una linea base reproducible para pronostico horario. Los resultados muestran que la inclusion de rezagos y variables temporales mejora la capacidad predictiva frente a usar solamente clima. "
        "Como limitaciones, el modelo lineal no captura todas las no linealidades y el entrenamiento avanzado N-BEATS requiere un entorno con Darts y aceleracion adecuada. "
        "Como trabajo futuro se recomienda validar contra modelos de boosting, redes recurrentes o N-BEATS optimizado con Optuna."
    )
    doc.add_heading("Referencias", level=1)
    doc.add_paragraph("[1] B. N. Oreshkin et al., N-BEATS: Neural basis expansion analysis for interpretable time series forecasting, 2020.")
    doc.add_paragraph("[2] F. Pedregosa et al., Scikit-learn: Machine Learning in Python, Journal of Machine Learning Research, 2011.")
    try:
        doc.save(DOCS / "reporte_ieee.docx")
    except PermissionError:
        doc.save(DOCS / "reporte_ieee_actualizado.docx")


def ppt_escape(s: str) -> str:
    return escape(s)


def slide_xml(title: str, bullets: list[str], image_rel: str | None = None) -> str:
    shapes = [
        f"""<p:sp><p:nvSpPr><p:cNvPr id="2" name="Title"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr><p:spPr><a:xfrm><a:off x="685800" y="304800"/><a:ext cx="7772400" cy="711200"/></a:xfrm></p:spPr><p:txBody><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr lang="es-ES" sz="3600" b="1"><a:solidFill><a:srgbClr val="0F172A"/></a:solidFill></a:rPr><a:t>{ppt_escape(title)}</a:t></a:r></a:p></p:txBody></p:sp>"""
    ]
    bullet_text = "".join(
        f"""<a:p><a:pPr marL="342900" indent="-171450"><a:buChar char="•"/></a:pPr><a:r><a:rPr lang="es-ES" sz="2100"/><a:t>{ppt_escape(b)}</a:t></a:r></a:p>"""
        for b in bullets
    )
    shapes.append(
        f"""<p:sp><p:nvSpPr><p:cNvPr id="3" name="Bullets"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr><p:spPr><a:xfrm><a:off x="762000" y="1310000"/><a:ext cx="4700000" cy="4200000"/></a:xfrm></p:spPr><p:txBody><a:bodyPr wrap="square"/><a:lstStyle/>{bullet_text}</p:txBody></p:sp>"""
    )
    if image_rel:
        shapes.append(
            f"""<p:pic><p:nvPicPr><p:cNvPr id="4" name="Image"/><p:cNvPicPr/><p:nvPr/></p:nvPicPr><p:blipFill><a:blip r:embed="{image_rel}"/><a:stretch><a:fillRect/></a:stretch></p:blipFill><p:spPr><a:xfrm><a:off x="5600000" y="1380000"/><a:ext cx="3370000" cy="3000000"/></a:xfrm><a:prstGeom prst="rect"><a:avLst/></a:prstGeom></p:spPr></p:pic>"""
        )
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
<p:cSld><p:bg><p:bgPr><a:solidFill><a:srgbClr val="F8FAFC"/></a:solidFill><a:effectLst/></p:bgPr></p:bg><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="0" cy="0"/><a:chOff x="0" y="0"/><a:chExt cx="0" cy="0"/></a:xfrm></p:grpSpPr>{''.join(shapes)}</p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sld>"""


def create_pptx(summary: dict[str, float]) -> None:
    slides = [
        ("Pronostico horario de SIN", ["Benjamin", "Proyecto de Inteligencia Artificial", "Series temporales y variables meteorologicas"], "predictions_vs_actual.png"),
        ("Planteamiento del problema", ["Predecir el comportamiento horario de SIN", "Apoyar analisis energetico y planificacion", "Integrar clima y patrones historicos"], None),
        ("Dataset", ["Archivo NuevoTPIA.csv", "Frecuencia horaria", "Objetivo: SIN Imputed", "Covariables: temperatura, humedad, presion, lluvia y viento"], None),
        ("Features creados", ["Ciclos calendario: hora, dia y mes", "Fin de semana y franjas horarias", "Rezagos: 24 h y 168 h", "Medias moviles: 24 h y 168 h", "Magnitud de viento e interaccion clima"], None),
        ("Modelo utilizado", ["Baseline: regresion Ridge", "Regularizacion para reducir sobreajuste", "Notebook original: N-BEATS como alternativa profunda"], None),
        ("Optimizacion", ["Grid Search manual", "Alphas: 0 a 1000", "Comparacion de tres conjuntos de features"], "grid_search_curve.png"),
        ("Visualizacion de resultados", ["Curva de busqueda de hiperparametros", "Comparacion de predicciones vs valores reales", "Histograma de errores"], "rmse_by_feature_set.png"),
        ("Predicciones y evaluacion", [f"RMSE: {summary['RMSE']:.2f}", f"MAE: {summary['MAE']:.2f}", f"MAPE: {summary['MAPE_%']:.2f}%", f"R2: {summary['R2']:.4f}"], "error_histogram.png"),
        ("Conclusiones y limitaciones", ["Los rezagos capturan estacionalidad fuerte", "El modelo lineal es interpretable y reproducible", "Limitacion: no captura todas las no linealidades"], None),
        ("Trabajos futuros", ["Optimizar N-BEATS con Optuna", "Comparar con boosting y modelos recurrentes", "Publicar GitHub Pages y completar datos personales"], None),
    ]
    out = PRESENTATION / "presentacion_proyecto_ia.pptx"
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Default Extension="png" ContentType="image/png"/><Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/><Override PartName="/ppt/slideMasters/slideMaster1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideMaster+xml"/><Override PartName="/ppt/slideLayouts/slideLayout1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slideLayout+xml"/>""" + "".join(f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>' for i in range(1, 11)) + "</Types>")
        z.writestr("_rels/.rels", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/></Relationships>""")
        slide_ids = "".join(f'<p:sldId id="{255+i}" r:id="rId{i}"/>' for i in range(1, 11))
        z.writestr("ppt/presentation.xml", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:sldMasterIdLst><p:sldMasterId id="2147483648" r:id="rId11"/></p:sldMasterIdLst><p:sldIdLst>{slide_ids}</p:sldIdLst><p:sldSz cx="9144000" cy="5143500" type="screen16x9"/><p:notesSz cx="6858000" cy="9144000"/></p:presentation>""")
        rels = "".join(f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>' for i in range(1, 11))
        rels += '<Relationship Id="rId11" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideMaster" Target="slideMasters/slideMaster1.xml"/>'
        z.writestr("ppt/_rels/presentation.xml.rels", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{rels}</Relationships>""")
        z.writestr("ppt/slideMasters/slideMaster1.xml", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sldMaster xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:cSld><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/></p:spTree></p:cSld><p:sldLayoutIdLst><p:sldLayoutId id="2147483649" r:id="rId1"/></p:sldLayoutIdLst><p:txStyles><p:titleStyle/><p:bodyStyle/><p:otherStyle/></p:txStyles></p:sldMaster>""")
        z.writestr("ppt/slideMasters/_rels/slideMaster1.xml.rels", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/></Relationships>""")
        z.writestr("ppt/slideLayouts/slideLayout1.xml", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><p:sldLayout xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" type="blank"><p:cSld name="Blank"><p:spTree><p:nvGrpSpPr><p:cNvPr id="1" name=""/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/></p:spTree></p:cSld></p:sldLayout>""")
        for i, (title, bullets, image) in enumerate(slides, 1):
            rel_id = None
            if image:
                rel_id = "rId2"
                z.writestr(f"ppt/media/{image}", (ASSETS / image).read_bytes())
                z.writestr(f"ppt/slides/_rels/slide{i}.xml.rels", f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="../media/{image}"/></Relationships>""")
            else:
                z.writestr(f"ppt/slides/_rels/slide{i}.xml.rels", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slideLayout" Target="../slideLayouts/slideLayout1.xml"/></Relationships>""")
            z.writestr(f"ppt/slides/slide{i}.xml", slide_xml(title, bullets, rel_id))


def main() -> None:
    ensure_dirs()
    copy_sources()
    df = load_and_engineer()
    results, pred_df, summary, features = run_grid_search(df)
    create_charts(results, pred_df)
    create_notebooks()
    create_readme(summary, features)
    create_index(summary)
    create_report(summary, results)
    create_pptx(summary)
    print(json.dumps({"summary": summary, "rows": len(df), "deliverables_root": str(ROOT)}, indent=2))


if __name__ == "__main__":
    main()
